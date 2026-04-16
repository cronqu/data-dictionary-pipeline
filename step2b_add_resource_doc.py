#!/usr/bin/env python3
"""
STEP 2b — Attach a resource document to inform data label definitions  (AUTOMATED)
=============================================================================
PURPOSE : Reads a user-supplied reference document (any format) and extracts
          its plain text content into resource_doc.txt — a format that can be
          appended to ai_input.txt so the AI tool uses it when writing
          variable descriptions.

SUPPORTED FORMATS:
  • Plain text   (.txt)
  • Markdown     (.md)
  • Word          (.docx)   — requires: pip3 install python-docx
  • Excel        (.xlsx)    — requires: pip3 install openpyxl
  • CSV          (.csv)
  • PDF          (.pdf)     — requires: pip3 install pypdf   OR  pdfminer.six

  NOTE: The resource document does NOT need to be any specific format.
        Any of the above formats will work. If your document is in a format
        not listed here (e.g., .odt, .rtf), convert it to .docx or .txt
        first using your word processor's "Save As" function.

OUTPUT  : resource_doc.txt  (plain text extract; reviewed by you before use)

PRIVACY : This step processes a REFERENCE document (e.g., a data dictionary,
          codebook, or annotation guide) — NOT your actual patient data CSV.
          The resource document should contain only variable definitions,
          code labels, and clinical descriptions — no patient records.
=============================================================================

EXAMPLES of suitable resource documents:
  • A data dictionary (Word, Excel, or PDF) from your data provider
  • A codebook describing variable labels and values
  • A clinical annotation guide (e.g., PARIS categories document, RAI manual)
  • Any document containing variable descriptions, code definitions, or
    clinical context that the AI should reference when writing descriptions

USAGE:
  python3 step2b_add_resource_doc.py path/to/your_reference_document.docx
  python3 step2b_add_resource_doc.py path/to/codebook.xlsx --sheet "Sheet1"
  python3 step2b_add_resource_doc.py path/to/data_guide.pdf
  python3 step2b_add_resource_doc.py path/to/labels.csv
  python3 step2b_add_resource_doc.py path/to/notes.txt
=============================================================================
"""

import os
import sys
import argparse

SCRIPT_DIR   = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE  = os.path.join(SCRIPT_DIR, "resource_doc.txt")

# ── Argument parsing ──────────────────────────────────────────────────────────
parser = argparse.ArgumentParser(
    description="Extract plain text from a reference document for AI context.",
    formatter_class=argparse.RawDescriptionHelpFormatter,
    epilog=__doc__,
)
parser.add_argument("document", help="Path to your reference document")
parser.add_argument("--sheet", default=None,
                    help="Sheet name (for .xlsx files; default: first sheet)")
parser.add_argument("--max-chars", type=int, default=60000,
                    help="Maximum characters to extract (default: 60000)")

args = parser.parse_args()
doc_path = args.document
max_chars = args.max_chars

if not os.path.exists(doc_path):
    print(f"ERROR: File not found: {doc_path}")
    sys.exit(1)

ext = os.path.splitext(doc_path)[1].lower()

# ── Extraction by file type ───────────────────────────────────────────────────
extracted_text = ""

if ext in (".txt", ".md"):
    # ── Plain text / Markdown ──
    with open(doc_path, "r", encoding="utf-8", errors="replace") as f:
        extracted_text = f.read()
    print(f"Read plain text from: {doc_path}")

elif ext == ".csv":
    # ── CSV ──
    import csv
    lines = []
    with open(doc_path, "r", encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            lines.append("  |  ".join(str(c) for c in row))
    extracted_text = "\n".join(lines)
    print(f"Read CSV from: {doc_path} ({len(lines)} rows)")

elif ext == ".docx":
    # ── Word document ──
    try:
        import docx as python_docx
    except ImportError:
        print("ERROR: python-docx is not installed.")
        print("Install it with:  pip3 install python-docx")
        sys.exit(1)

    doc = python_docx.Document(doc_path)
    parts = []

    # Main paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables
    for table in doc.tables:
        parts.append("")
        for row in table.rows:
            parts.append("  |  ".join(c.text.strip() for c in row.cells))

    # Comments (if any — useful for annotation guides like the PARIS document)
    try:
        import zipfile, xml.etree.ElementTree as ET
        with zipfile.ZipFile(doc_path) as z:
            if "word/comments.xml" in z.namelist():
                parts.append("\n── DOCUMENT COMMENTS ──")
                with z.open("word/comments.xml") as f:
                    tree = ET.parse(f)
                    root = tree.getroot()
                    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                    for comment in root.findall(".//w:comment", ns):
                        author = comment.get(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author",
                            "Unknown")
                        texts = [t.text for t in comment.iter(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                                 if t.text]
                        text = "".join(texts)
                        if text.strip():
                            parts.append(f"[Comment by {author}]: {text}")
        print("  (also extracted embedded comments from .docx)")
    except Exception:
        pass  # Comments optional

    extracted_text = "\n".join(parts)
    print(f"Read Word document from: {doc_path}")

elif ext == ".xlsx":
    # ── Excel ──
    try:
        import openpyxl
    except ImportError:
        print("ERROR: openpyxl is not installed.")
        print("Install it with:  pip3 install openpyxl")
        sys.exit(1)

    wb = openpyxl.load_workbook(doc_path, data_only=True)
    target_sheet = args.sheet or wb.sheetnames[0]
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"\n── Sheet: {sheet_name} ──")
        for row in ws.iter_rows(values_only=True):
            row_vals = [str(c) if c is not None else "" for c in row]
            if any(v.strip() for v in row_vals):
                parts.append("  |  ".join(row_vals))
    extracted_text = "\n".join(parts)
    print(f"Read Excel from: {doc_path} ({len(wb.sheetnames)} sheets)")

elif ext == ".pdf":
    # ── PDF — try pypdf first, then pdfminer ──
    extracted_text = ""
    try:
        from pypdf import PdfReader
        reader = PdfReader(doc_path)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append(f"── Page {i+1} ──\n{text}")
        extracted_text = "\n\n".join(pages)
        print(f"Read PDF with pypdf from: {doc_path} ({len(reader.pages)} pages)")
    except ImportError:
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            extracted_text = pdfminer_extract(doc_path)
            print(f"Read PDF with pdfminer from: {doc_path}")
        except ImportError:
            print("ERROR: No PDF library found.")
            print("Install one of:")
            print("  pip3 install pypdf")
            print("  pip3 install pdfminer.six")
            sys.exit(1)

else:
    print(f"ERROR: Unsupported file format: {ext}")
    print()
    print("Supported formats: .txt  .md  .csv  .docx  .xlsx  .pdf")
    print()
    print("If your file is in another format (e.g., .odt, .rtf),")
    print("convert it to .docx or .txt using your word processor's 'Save As'.")
    sys.exit(1)

# ── Truncate if needed ────────────────────────────────────────────────────────
char_count = len(extracted_text)
if char_count > max_chars:
    extracted_text = extracted_text[:max_chars]
    print(f"WARNING: Document truncated to {max_chars:,} characters "
          f"(original: {char_count:,}). Use --max-chars N to increase the limit.")

# ── Write resource_doc.txt ────────────────────────────────────────────────────
header = (
    f"RESOURCE DOCUMENT: {os.path.basename(doc_path)}\n"
    f"{'=' * 70}\n"
    f"This document was provided by the user as a reference for defining\n"
    f"data variable labels and descriptions. The AI should use it to inform\n"
    f"the 'description' field of each variable in the data dictionary.\n"
    f"{'=' * 70}\n\n"
)

with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
    f.write(header + extracted_text)

print(f"✅  Extracted {char_count:,} characters → {OUTPUT_FILE}")
print()
print("──────────────────────────────────────────────────────────────────")
print("REVIEW (recommended):")
print("──────────────────────────────────────────────────────────────────")
print("Open resource_doc.txt and confirm:")
print("  • It contains variable definitions / code labels (not patient data)")
print("  • Key sections relevant to your column names are present")
print("  • You are comfortable with this content being sent to the AI tool")
print()
print("NEXT STEP:")
print("  Run:  python3 step3_package_for_ai.py")
print("  (resource_doc.txt will be automatically included in ai_input.txt)")
print()
